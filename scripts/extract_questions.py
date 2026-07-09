from __future__ import annotations

import hashlib
import json
import re
import unicodedata
from pathlib import Path
from typing import Any

from docx import Document

PROJECT = Path(__file__).resolve().parents[1]
SOURCE = PROJECT.parents[2] if PROJECT.parent.name == ".worktrees" else PROJECT.parent
DATA = PROJECT / "data"
YXQ_DOCUMENT = SOURCE / "yxq英语题库.docx"
HR_DOCUMENT = SOURCE / "人力期中考核总复习_修正版答案解析.docx"
ENGLISH_REVIEW_DOCUMENT = SOURCE / "英语复习.docx"

LETTERS = "ABCDEFGHIJKLMNO"
OPTION_KEYS = [*LETTERS, "T"]


def normalize(value: str) -> str:
    value = unicodedata.normalize("NFKC", value).lower()
    value = re.sub(r"[^\w\s]", " ", value, flags=re.UNICODE)
    return re.sub(r"\s+", " ", value).strip()


def fingerprint(question: str, options: list[dict[str, str]], salt: str = "") -> str:
    canonical = "|".join([salt, normalize(question)] + [f"{o['key']}:{normalize(o['text'])}" for o in options])
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def source_label(path: Path) -> str:
    name = path.name
    if name == "英语复习.docx":
        return "英语复习.docx"
    if path.stem.lower().startswith("yxq"):
        return "yxq英语题库.docx"
    if "人力" in name:
        return "人力期中考核总复习_修正版答案解析.docx"
    return name


def subject_for_document(path: Path) -> str:
    return "human_resources" if "人力" in source_label(path) else "english"


def category_for_document(path: Path) -> str:
    if source_label(path) == "英语复习.docx":
        return "english_review"
    return "human_resources_review" if subject_for_document(path) == "human_resources" else "vocabulary"


def structured_documents() -> list[Path]:
    docs: list[Path] = []
    if YXQ_DOCUMENT.exists():
        docs.append(YXQ_DOCUMENT)
    if HR_DOCUMENT.exists():
        docs.append(HR_DOCUMENT)
    if ENGLISH_REVIEW_DOCUMENT.exists():
        docs.append(ENGLISH_REVIEW_DOCUMENT)
    if docs:
        return docs

    for path in SOURCE.glob("*.docx"):
        if path.name.startswith("~$"):
            continue
        name = path.name.lower()
        if "yxq" in name or "人力" in path.name or path.name == "英语复习.docx":
            docs.append(path)
    return sorted(docs, key=lambda p: p.name)


def make_record(
    *,
    path: Path,
    number: str,
    question_type: str,
    question: str,
    options: list[dict[str, str]],
    answer: list[str],
    explanation: str = "",
    answer_text: str = "",
) -> dict[str, Any]:
    subject = subject_for_document(path)
    salt = f"{subject}:{number}" if subject == "human_resources" else ""
    fp = fingerprint(question, options, salt)
    option_keys = {option["key"] for option in options}
    if question_type == "short_answer":
        complete = bool(question.strip()) and bool((answer_text or explanation).strip())
    else:
        complete = bool(question.strip()) and len(options) >= 2 and bool(answer) and all(item in option_keys for item in answer)

    record: dict[str, Any] = {
        "id": f"{'hr' if subject == 'human_resources' else 'q'}_{fp[:16]}",
        "subject": subject,
        "type": question_type,
        "question": question.strip(),
        "options": options,
        "answer": answer,
        "category": category_for_document(path),
        "source": [f"{source_label(path)}#{number}"],
        "explanation": explanation.strip(),
        "ocrConfidence": None,
        "reviewStatus": "verified" if complete else "needs_review",
        "fingerprint": fp,
    }
    if answer_text.strip():
        record["answerText"] = answer_text.strip()
    return record


def parse_answer_letters(text: str) -> list[str]:
    return [letter for letter in re.findall(r"[A-DTF]", text.upper())]


def parse_english_document(path: Path) -> list[dict[str, Any]]:
    lines = [p.text.strip() for p in Document(path).paragraphs if p.text.strip()]
    starts = [i for i, line in enumerate(lines) if re.match(r"^\d+\.(?:\s+.*)?$", line)]
    records = []
    for position, start in enumerate(starts):
        end = starts[position + 1] if position + 1 < len(starts) else len(lines)
        match = re.match(r"^(\d+)\.(?:\s+(.*))?$", lines[start])
        assert match
        number = match.group(1)
        question_parts = [match.group(2).strip()] if match.group(2) else []
        options: list[dict[str, str]] = []
        marked_answer = None
        stated_answer = None
        for line in lines[start + 1:end]:
            answer_match = re.search(r"[：:]\s*([A-DTF])\s*$", line)
            if line.startswith(("正确答案", "答案")) and answer_match:
                stated_answer = answer_match.group(1)
                continue
            option_match = re.match(r"^(✓\s*)?([A-D])\.\s*(.+)$", line)
            if option_match:
                if option_match.group(1):
                    marked_answer = option_match.group(2)
                options.append({"key": option_match.group(2), "text": option_match.group(3).strip()})
                continue
            if not options:
                question_parts.append(line)
        answer = stated_answer or marked_answer
        question = " ".join(question_parts).strip()
        question_type = "single_choice"
        if not options and answer in {"T", "F"}:
            question_type = "true_false"
            options = [{"key": "T", "text": "正确"}, {"key": "F", "text": "错误"}]
        records.append(make_record(
            path=path,
            number=number,
            question_type=question_type,
            question=question,
            options=options,
            answer=[answer] if answer else [],
        ))
    return records


def option_match(line: str) -> re.Match[str] | None:
    return re.match(r"^([A-D])\.\s*(.+)$", line)


def objective_start(line: str) -> re.Match[str] | None:
    return re.match(r"^(\d+)\.\s*(.+)$", line) or re.match(r"^(案例[（(][一二三四五六七八九十]+[）)]-\d+)\.\s*(.+)$", line)


def parse_objective_block(path: Path, number: str, question: str, body: list[str], question_type: str) -> dict[str, Any]:
    options: list[dict[str, str]] = []
    answer: list[str] = []
    explanation = ""
    for line in body:
        match = option_match(line)
        if match:
            options.append({"key": match.group(1), "text": match.group(2).strip()})
            continue
        if line.startswith("答案"):
            answer = parse_answer_letters(line.split("：", 1)[-1].split(":", 1)[-1])
            continue
        if line.startswith("解析"):
            explanation = line.split("：", 1)[-1].split(":", 1)[-1].strip()
    return make_record(path=path, number=number, question_type=question_type, question=question, options=options, answer=answer, explanation=explanation)


def parse_short_block(path: Path, number: str, question: str, body: list[str]) -> dict[str, Any]:
    answer_text = ""
    explanation = ""
    for line in body:
        if line.startswith("答案") or line.startswith("参考答案"):
            answer_text = line.split("：", 1)[-1].split(":", 1)[-1].strip()
        elif line.startswith("解析"):
            explanation = line.split("：", 1)[-1].split(":", 1)[-1].strip()
        else:
            question = f"{question}\n{line}".strip()
    return make_record(path=path, number=number, question_type="short_answer", question=question, options=[], answer=[], explanation=explanation, answer_text=answer_text)


def parse_hr_document(path: Path) -> list[dict[str, Any]]:
    lines = [p.text.strip() for p in Document(path).paragraphs if p.text.strip()]
    records: list[dict[str, Any]] = []
    section = ""
    active: dict[str, Any] | None = None
    case_contexts: dict[str, str] = {}
    current_case_context = ""

    def flush() -> None:
        nonlocal active
        if not active:
            return
        if active["type"] in {"single_choice", "multiple_choice"}:
            records.append(parse_objective_block(path, active["number"], active["question"], active["body"], active["type"]))
        else:
            records.append(parse_short_block(path, active["number"], active["question"], active["body"]))
        active = None

    for line in lines:
        if line.startswith("一、"):
            flush()
            section = "single_choice"
            continue
        if line.startswith("二、"):
            flush()
            section = "multiple_choice"
            continue
        if line.startswith("三、"):
            flush()
            section = "calculation"
            current_case_context = ""
            continue
        if line.startswith("四、"):
            flush()
            section = "case_analysis"
            current_case_context = ""
            continue
        if not section:
            continue

        context_match = re.match(r"^(案例[（(][一二三四五六七八九十]+[）)]).+", line)
        start = objective_start(line)
        if section == "multiple_choice" and context_match and not start:
            case_contexts[context_match.group(1)] = line
            continue
        if section == "case_analysis" and not re.match(r"^\d+\.", line) and not line.startswith(("答案", "参考答案", "解析")):
            current_case_context = line
            continue

        if section in {"single_choice", "multiple_choice"} and start:
            flush()
            number = start.group(1)
            question = start.group(2)
            case_key_match = re.match(r"^(案例[（(][一二三四五六七八九十]+[）)])", number)
            if case_key_match and case_key_match.group(1) in case_contexts:
                question = f"{case_contexts[case_key_match.group(1)]}\n{number}. {question}"
            active = {"number": number, "question": question, "body": [], "type": section}
            continue

        short_match = re.match(r"^(\d+)\.\s*(.+)$", line)
        if section in {"calculation", "case_analysis"} and short_match:
            flush()
            question = short_match.group(2)
            if section == "case_analysis" and current_case_context:
                question = f"{current_case_context}\n{question}"
            active = {"number": f"{section}-{short_match.group(1)}", "question": question, "body": [], "type": "short_answer"}
            continue

        if active:
            active["body"].append(line)

    flush()
    return records


def parse_inline_choice(text: str) -> tuple[str, list[str]] | None:
    match = re.search(r"\(([^()]+(?:/|,)[^()]+)\)", text)
    if not match:
        return None
    parts = [part.strip().lower() for part in re.split(r"\s*(?:/|,)\s*", match.group(1)) if part.strip()]
    return (match.group(0), parts) if len(parts) == 2 else None


def key_for_answer(options: list[dict[str, str]], answer_text: str) -> str:
    answer_norm = normalize(answer_text)
    for option in options:
        if normalize(option["text"]) == answer_norm:
            return option["key"]
    return ""


def parse_phrase_section(path: Path, lines: list[str]) -> list[dict[str, Any]]:
    numbered = [(index, re.match(r"^(\d+)\.\s*(.+)$", line)) for index, line in enumerate(lines) if re.match(r"^\d+\.\s*", line)]
    if not numbered:
        return []
    last_numbered_index = numbered[-1][0]
    answers = [line.strip() for line in lines[last_numbered_index + 1:] if line.strip()]
    records: list[dict[str, Any]] = []
    if len(answers) != len(numbered):
        for index, match in numbered:
            assert match
            records.append(make_record(path=path, number=f"part1-phrase-{match.group(1)}", question_type="single_choice", question=match.group(2), options=[], answer=[]))
        return records

    options = [{"key": OPTION_KEYS[index], "text": answer} for index, answer in enumerate(answers)]
    for position, (line_index, match) in enumerate(numbered):
        assert match
        next_index = numbered[position + 1][0] if position + 1 < len(numbered) else last_numbered_index + 1
        parts = [match.group(2), *lines[line_index + 1:next_index]]
        answer_key = options[position]["key"]
        records.append(make_record(
            path=path,
            number=f"part1-phrase-{match.group(1)}",
            question_type="single_choice",
            question=" ".join(parts),
            options=options,
            answer=[answer_key],
        ))
    return records


def grouped_choice_indices(lines: list[str]) -> list[list[int]]:
    indices: list[tuple[int, int]] = []
    for index, line in enumerate(lines):
        match = re.match(r"^(\d+)\.\s*(.+)$", line)
        if match and parse_inline_choice(line):
            indices.append((index, int(match.group(1))))
    groups: list[list[int]] = []
    current: list[int] = []
    previous_number = 0
    for index, number in indices:
        if current and number <= previous_number:
            groups.append(current)
            current = []
        current.append(index)
        previous_number = number
    if current:
        groups.append(current)
    return groups


def parse_choice_sections(path: Path, lines: list[str]) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    groups = grouped_choice_indices(lines)
    for group_number, starts in enumerate(groups, start=1):
        group_end = groups[group_number][0] if group_number < len(groups) else len(lines)
        answer_start = max(starts[-1] + 1, group_end - len(starts))
        answers = [line.strip() for line in lines[answer_start:group_end] if line.strip()]
        for position, start in enumerate(starts):
            line = lines[start]
            match = re.match(r"^(\d+)\.\s*(.+)$", line)
            assert match
            next_start = starts[position + 1] if position + 1 < len(starts) else answer_start
            question = " ".join([match.group(2), *lines[start + 1:next_start]])
            parsed = parse_inline_choice(question)
            if not parsed:
                continue
            marker, choices = parsed
            options = [{"key": "A", "text": choices[0]}, {"key": "B", "text": choices[1]}]
            answer_text = answers[position] if position < len(answers) else ""
            answer = key_for_answer(options, answer_text)
            records.append(make_record(
                path=path,
                number=f"part1-choice-{len(records) + 1}",
                question_type="single_choice",
                question=question.replace(marker, "______").strip(),
                options=options,
                answer=[answer] if answer else [],
            ))
    return records


def option_pairs(line: str) -> list[tuple[str, str]]:
    pairs = []
    for key, word in re.findall(r"\b([A-O0W])\s+([A-Za-z][A-Za-z-]*)", line):
        fixed_key = "O" if key == "0" else key
        pairs.append((fixed_key, word))
    return pairs


def normalize_bank(raw_pairs: list[tuple[str, str]]) -> list[dict[str, str]]:
    keys = {key for key, _ in raw_pairs}
    options: list[dict[str, str]] = []
    seen: set[str] = set()
    for key, word in raw_pairs:
        fixed_key = "M" if key == "W" and "M" not in keys else key
        if fixed_key in seen or fixed_key not in LETTERS:
            continue
        seen.add(fixed_key)
        options.append({"key": fixed_key, "text": word})
    return sorted(options, key=lambda option: LETTERS.index(option["key"]))


def parse_answer_mapping(lines: list[str]) -> dict[int, str]:
    mapping: dict[int, str] = {}
    for line in lines:
        for start, end, letters in re.findall(r"(?:(\d+)\s*)?[—-]+\s*(\d+)\s*([A-Z0-9]+)", line.upper()):
            end_number = int(end)
            start_number = int(start) if start else (1 if end_number == 5 else end_number - len(letters) + 1)
            for offset, letter in enumerate(letters):
                mapping[start_number + offset] = "O" if letter == "0" else letter
    return mapping


def is_mapping_line(line: str) -> bool:
    return bool(re.search(r"[—-]+\s*(?:5|10)\s*[A-Z0-9]+", line.upper()))


def parse_cloze_sections(path: Path, lines: list[str]) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    passage = 0
    index = 0
    while index < len(lines):
        raw_pairs = option_pairs(lines[index])
        if not raw_pairs:
            index += 1
            continue
        bank_pairs = raw_pairs[:]
        index += 1
        while index < len(lines) and option_pairs(lines[index]):
            bank_pairs.extend(option_pairs(lines[index]))
            index += 1
        body: list[str] = []
        answer_lines: list[str] = []
        while index < len(lines) and not option_pairs(lines[index]):
            if is_mapping_line(lines[index]):
                answer_lines.append(lines[index])
            else:
                body.append(lines[index])
            index += 1
        passage += 1
        options = normalize_bank(bank_pairs)
        answers = parse_answer_mapping(answer_lines)
        context = " ".join(body).strip()
        for blank in sorted(answers):
            records.append(make_record(
                path=path,
                number=f"part2-passage{passage}-blank{blank}",
                question_type="single_choice",
                question=f"Passage {passage} blank ({blank}): {context}",
                options=options,
                answer=[answers[blank]],
            ))
    return records


def parse_english_review_document(path: Path) -> list[dict[str, Any]]:
    lines = [p.text.strip() for p in Document(path).paragraphs if p.text.strip()]
    part2_index = next((index for index, line in enumerate(lines) if line.startswith("第二部分")), len(lines))
    part1 = [line for line in lines[:part2_index] if not line.startswith(("英语复习", "第一部分"))]
    part2 = lines[part2_index + 1:] if part2_index < len(lines) else []
    first_choice = next((index for index, line in enumerate(part1) if parse_inline_choice(line)), len(part1))
    return [
        *parse_phrase_section(path, part1[:first_choice]),
        *parse_choice_sections(path, part1[first_choice:]),
        *parse_cloze_sections(path, part2),
    ]


def parse_document(path: Path) -> list[dict[str, Any]]:
    if source_label(path) == "英语复习.docx":
        return parse_english_review_document(path)
    if subject_for_document(path) == "human_resources":
        return parse_hr_document(path)
    return parse_english_document(path)


def count_ocr_pages() -> tuple[int, int]:
    files = list((SOURCE / "tmp" / "ocr").glob("*/ocr.json"))
    pages = 0
    for path in files:
        text = path.read_text(encoding="utf-8", errors="replace")
        pages += len(re.findall(r'"index"\s*:', text))
    return len(files), pages


def main() -> None:
    DATA.mkdir(parents=True, exist_ok=True)
    docs = structured_documents()
    records = [record for path in docs for record in parse_document(path)]
    ocr_files, ocr_pages = count_ocr_pages()
    subject_counts: dict[str, int] = {}
    type_counts: dict[str, int] = {}
    category_counts: dict[str, int] = {}
    for record in records:
        subject_counts[record["subject"]] = subject_counts.get(record["subject"], 0) + 1
        type_counts[f"{record['subject']}:{record['type']}"] = type_counts.get(f"{record['subject']}:{record['type']}", 0) + 1
        category_counts[f"{record['subject']}:{record['category']}"] = category_counts.get(f"{record['subject']}:{record['category']}", 0) + 1
    (DATA / "questions.raw.json").write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
    (DATA / "extraction-meta.json").write_text(json.dumps({
        "sourceDocuments": [source_label(p) for p in docs],
        "structuredDocuments": [source_label(p) for p in docs],
        "structuredRecognized": len(records),
        "subjectCounts": subject_counts,
        "typeCounts": type_counts,
        "categoryCounts": category_counts,
        "ocrFilesInspected": ocr_files,
        "ocrPagesInspected": ocr_pages,
        "ocrCandidatesAccepted": 0,
        "ocrPolicy": "OCR pages are retained for later manual segmentation; no ambiguous OCR candidate is promoted automatically."
    }, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"structured": len(records), "subjects": subject_counts, "types": type_counts, "categories": category_counts, "ocr_pages": ocr_pages}, ensure_ascii=False))


if __name__ == "__main__":
    main()
